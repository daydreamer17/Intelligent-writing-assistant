from __future__ import annotations

from io import BytesIO
import re
import json
from uuid import UUID, uuid4

import logging

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile

from ...models.schemas import (
    DeleteDocumentResponse,
    SearchDocumentsRequest,
    SearchDocumentsResponse,
    SourceDocumentResponse,
    UploadDocumentsRequest,
    UploadDocumentsResponse,
    RetrievalEvalRequest,
    RetrievalEvalResponse,
    RetrievalEvalCaseResult,
    RetrievalMetricAtK,
    RetrievalEvalRunsResponse,
    RetrievalEvalRunSummaryResponse,
    DeleteRetrievalEvalRunResponse,
)
from ...services.rag_service import UploadedDocument
from ...services.retrieval_eval_service import EvalCase, RetrievalEvalService
from ..deps import AppServices, get_services

router = APIRouter(tags=["rag"])
logger = logging.getLogger("app.rag")
_eval_service = RetrievalEvalService()


@router.post("/rag/upload", response_model=UploadDocumentsResponse)
def upload_documents(
    payload: UploadDocumentsRequest,
    services: AppServices = Depends(get_services),
) -> UploadDocumentsResponse:
    try:
        max_bytes = services.upload_max_bytes
        docs = [
            UploadedDocument(
                doc_id=_normalize_doc_id(item.doc_id),
                title=item.title,
                content=item.content,
                url=item.url,
            )
            for item in payload.documents
        ]
        for doc in docs:
            if len(doc.content.encode("utf-8", errors="ignore")) > max_bytes:
                raise HTTPException(status_code=413, detail="Document too large")
        added = services.rag.add_documents(docs)
        response_docs = [
            SourceDocumentResponse(
                doc_id=doc.doc_id,
                title=doc.title,
                content=doc.content,
                url=doc.url,
            )
            for doc in added
        ]
        return UploadDocumentsResponse(documents=response_docs)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG upload failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/rag/search", response_model=SearchDocumentsResponse)
def search_documents(
    payload: SearchDocumentsRequest,
    services: AppServices = Depends(get_services),
) -> SearchDocumentsResponse:
    try:
        results = services.rag.search(payload.query, top_k=payload.top_k)
        response_docs = [
            SourceDocumentResponse(
                doc_id=doc.doc_id,
                title=doc.title,
                content=doc.content,
                url=doc.url,
            )
            for doc in results
        ]
        return SearchDocumentsResponse(documents=response_docs)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG search failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/rag/evaluate", response_model=RetrievalEvalResponse)
def evaluate_retrieval(
    payload: RetrievalEvalRequest,
    services: AppServices = Depends(get_services),
) -> RetrievalEvalResponse:
    try:
        if not payload.cases:
            raise HTTPException(status_code=422, detail="No evaluation cases provided")

        k_values = [k for k in payload.k_values if k > 0] or [1, 3, 5]
        max_k = max(k_values)

        cases = [
            EvalCase(
                query=case.query,
                relevant_doc_ids=case.relevant_doc_ids,
                query_id=case.query_id,
            )
            for case in payload.cases
            if case.query.strip()
        ]
        if not cases:
            raise HTTPException(status_code=422, detail="All evaluation queries are empty")
        rag_override = payload.rag_config_override
        if rag_override is not None:
            logger.info(
                "RAG evaluate override received: rerank=%s hyde=%s bilingual=%s",
                rag_override.rerank_enabled,
                rag_override.hyde_enabled,
                rag_override.bilingual_rewrite_enabled,
            )

        report = _eval_service.evaluate(
            cases=cases,
            k_values=k_values,
            search_fn=lambda query, top_k: services.rag.search(
                query,
                top_k=min(top_k, max_k),
                rag_eval_override=rag_override,
            ),
        )

        macro_metrics_data = [_metric_to_dict(item) for item in report.macro_metrics]
        per_query_data = [
            {
                "query": row.query,
                "query_id": row.query_id,
                "relevant_count": row.relevant_count,
                "retrieved_doc_ids": row.retrieved_doc_ids,
                "metrics": [_metric_to_dict(m) for m in row.metrics],
            }
            for row in report.per_query
        ]
        run_id = services.storage.save_retrieval_eval_run(
            total_queries=report.total_queries,
            queries_with_relevance=report.queries_with_relevance,
            k_values=report.k_values,
            macro_metrics=macro_metrics_data,
            per_query=per_query_data,
        )
        saved_row = services.storage.get_retrieval_eval_run(run_id)

        return RetrievalEvalResponse(
            eval_run_id=run_id,
            created_at=saved_row.created_at if saved_row else "",
            total_queries=report.total_queries,
            queries_with_relevance=report.queries_with_relevance,
            k_values=report.k_values,
            macro_metrics=[_metric_from_dict(item) for item in macro_metrics_data],
            per_query=[_case_result_from_dict(item) for item in per_query_data],
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG evaluate failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.get("/rag/evaluations", response_model=RetrievalEvalRunsResponse)
def list_retrieval_evaluations(
    limit: int = 20,
    services: AppServices = Depends(get_services),
) -> RetrievalEvalRunsResponse:
    try:
        runs = services.storage.list_retrieval_eval_runs(limit=limit)
        response_runs = [
            RetrievalEvalRunSummaryResponse(
                run_id=row.run_id,
                created_at=row.created_at,
                total_queries=row.total_queries,
                queries_with_relevance=row.queries_with_relevance,
                k_values=_json_load_list(row.k_values_json, default=[]),
                macro_metrics=[
                    _metric_from_dict(item)
                    for item in _json_load_list(row.macro_metrics_json, default=[])
                ],
            )
            for row in runs
        ]
        return RetrievalEvalRunsResponse(runs=response_runs)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG evaluations list failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.get("/rag/evaluations/{run_id}", response_model=RetrievalEvalResponse)
def get_retrieval_evaluation(
    run_id: int,
    services: AppServices = Depends(get_services),
) -> RetrievalEvalResponse:
    try:
        row = services.storage.get_retrieval_eval_run(run_id)
        if row is None:
            raise HTTPException(status_code=404, detail="Evaluation run not found")
        macro_metrics_data = _json_load_list(row.macro_metrics_json, default=[])
        per_query_data = _json_load_list(row.per_query_json, default=[])
        return RetrievalEvalResponse(
            eval_run_id=row.run_id,
            created_at=row.created_at,
            total_queries=row.total_queries,
            queries_with_relevance=row.queries_with_relevance,
            k_values=_json_load_list(row.k_values_json, default=[]),
            macro_metrics=[_metric_from_dict(item) for item in macro_metrics_data],
            per_query=[_case_result_from_dict(item) for item in per_query_data],
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG evaluation detail failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.delete("/rag/evaluations/{run_id}", response_model=DeleteRetrievalEvalRunResponse)
def delete_retrieval_evaluation(
    run_id: int,
    services: AppServices = Depends(get_services),
) -> DeleteRetrievalEvalRunResponse:
    try:
        deleted = services.storage.delete_retrieval_eval_run(run_id)
        return DeleteRetrievalEvalRunResponse(deleted=deleted)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG evaluation delete failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc

@router.get("/rag/documents", response_model=SearchDocumentsResponse)
def list_documents(
    limit: int = 100,
    services: AppServices = Depends(get_services),
) -> SearchDocumentsResponse:
    try:
        docs = services.rag.list_documents()
        if limit:
            docs = docs[: max(1, limit)]
        response_docs = [
            SourceDocumentResponse(
                doc_id=doc.doc_id,
                title=doc.title,
                content=doc.content,
                url=doc.url,
            )
            for doc in docs
        ]
        return SearchDocumentsResponse(documents=response_docs)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG list failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.delete("/rag/documents/{doc_id}", response_model=DeleteDocumentResponse)
def delete_document(
    doc_id: str,
    services: AppServices = Depends(get_services),
) -> DeleteDocumentResponse:
    try:
        deleted = services.rag.delete_document(doc_id)
        return DeleteDocumentResponse(deleted=deleted)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG delete failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/rag/upload-file", response_model=UploadDocumentsResponse)
async def upload_files(
    files: list[UploadFile] | None = File(default=None, description="上传文本文件"),
    services: AppServices = Depends(get_services),
) -> UploadDocumentsResponse:
    try:
        if not files:
            raise HTTPException(status_code=422, detail="No files provided")
        max_bytes = services.upload_max_bytes
        documents: list[UploadedDocument] = []
        for file in files:
            content_bytes = await file.read()
            if len(content_bytes) > max_bytes:
                raise HTTPException(status_code=413, detail="File too large")
            content = _extract_text(file.filename or "", content_bytes)
            documents.append(
                UploadedDocument(
                    doc_id=str(uuid4()),
                    title=file.filename or "uploaded",
                    content=content,
                    url="",
                )
            )
        added = services.rag.add_documents(documents)
        response_docs = [
            SourceDocumentResponse(
                doc_id=doc.doc_id,
                title=doc.title,
                content=doc.content,
                url=doc.url,
            )
            for doc in added
        ]
        return UploadDocumentsResponse(documents=response_docs)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("RAG file upload failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


def _extract_text(filename: str, content_bytes: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(content_bytes)
    if name.endswith(".docx"):
        return _extract_docx_text(content_bytes)
    if name.endswith(".md") or name.endswith(".markdown"):
        return _extract_markdown_text(content_bytes)
    return content_bytes.decode("utf-8", errors="ignore")


def _normalize_doc_id(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return str(uuid4())
    if raw.isdigit():
        return raw
    try:
        UUID(raw)
        return raw
    except Exception:
        return str(uuid4())


def _extract_pdf_text(content_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise HTTPException(status_code=400, detail="PDF support not installed") from exc
    reader = PdfReader(BytesIO(content_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _extract_docx_text(content_bytes: bytes) -> str:
    try:
        import docx
    except Exception as exc:
        raise HTTPException(status_code=400, detail="DOCX support not installed") from exc
    file_like = BytesIO(content_bytes)
    document = docx.Document(file_like)
    parts = [para.text for para in document.paragraphs if para.text]
    return "\n".join(parts)


def _extract_markdown_text(content_bytes: bytes) -> str:
    text = content_bytes.decode("utf-8", errors="ignore")
    # Remove fenced code blocks
    text = re.sub(r"```.*?```", "", text, flags=re.S)
    # Remove inline code
    text = re.sub(r"`([^`]*)`", r"\1", text)
    # Convert links/images to their label
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Strip markdown headings and blockquotes
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"^\s{0,3}>\s*", "", text, flags=re.M)
    # Strip list markers
    text = re.sub(r"^\s*([-*+]|\d+\.)\s+", "", text, flags=re.M)
    # Remove emphasis markers
    text = re.sub(r"[*_]{1,3}", "", text)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def _metric_to_dict(item: object) -> dict:
    return {
        "k": int(getattr(item, "k")),
        "recall": float(getattr(item, "recall")),
        "precision": float(getattr(item, "precision")),
        "hit_rate": float(getattr(item, "hit_rate")),
        "mrr": float(getattr(item, "mrr")),
        "ndcg": float(getattr(item, "ndcg")),
    }


def _metric_from_dict(item: dict) -> RetrievalMetricAtK:
    return RetrievalMetricAtK(
        k=int(item.get("k", 0)),
        recall=float(item.get("recall", 0.0)),
        precision=float(item.get("precision", 0.0)),
        hit_rate=float(item.get("hit_rate", 0.0)),
        mrr=float(item.get("mrr", 0.0)),
        ndcg=float(item.get("ndcg", 0.0)),
    )


def _case_result_from_dict(item: dict) -> RetrievalEvalCaseResult:
    return RetrievalEvalCaseResult(
        query=str(item.get("query", "")),
        query_id=str(item.get("query_id", "")),
        relevant_count=int(item.get("relevant_count", 0)),
        retrieved_doc_ids=[str(v) for v in item.get("retrieved_doc_ids", [])],
        metrics=[_metric_from_dict(m) for m in item.get("metrics", [])],
    )


def _json_load_list(raw: str, default: list) -> list:
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return data
        return default
    except Exception:
        return default
